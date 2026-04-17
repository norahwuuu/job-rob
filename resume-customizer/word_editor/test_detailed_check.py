"""
详细检查修改日志
"""
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent))

from resume_modifier.ai_analyzer import AIAnalyzer
from resume_modifier.content_modifier import ContentModifier

def check_modifications():
    """检查修改结果和失败原因"""
    
    print("=" * 80)
    print("🔍 详细修改日志检查")
    print("=" * 80)
    
    resume_path = r"./cv.docx"
    
    # 读取简历
    doc = Document(resume_path)
    resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    # 简单测试JD
    job_description = """
    AI Engineer Position
    
    We are looking for an AI Engineer with experience in:
    - LLM API integrations and prompt engineering
    - RAG systems and semantic search
    - FastAPI backend development
    - Customer collaboration and prototyping
    - Evaluation and monitoring tools (e.g., Langfuse)
    
    The role requires strong Python skills and experience with LangChain.
    """
    
    print("\n🤖 AI 分析中...")
    analyzer = AIAnalyzer()
    result = analyzer.analyze(job_description, resume_text)
    
    print(f"\n生成了 {len(result.modifications)} 条修改指令\n")
    
    # 应用修改
    modifier = ContentModifier(resume_path)
    
    for i, mod in enumerate(result.modifications, 1):
        print(f"\n{'='*70}")
        print(f"修改 {i}/{len(result.modifications)}")
        print(f"{'='*70}")
        print(f"Target: {mod.target[:80]}...")
        print(f"Replacement: {mod.replacement[:80]}...")
        print(f"Match Type: {mod.match_type}")
        
        success = modifier._apply_single_modification(mod)
        
        if success:
            print(f"✅ 成功")
        else:
            print(f"❌ 失败")
            # 查找日志中的错误信息
            for log in modifier.logs:
                if log.target == mod.target and not log.success:
                    print(f"   错误: {log.error_message}")
                    if log.debug_info:
                        print(f"   调试信息: {log.debug_info}")
                    break
    
    print("\n" + "="*80)
    print(f"总结: 成功 {sum(1 for log in modifier.logs if log.success)}/{len(result.modifications)}")
    print("="*80)
    
    # 保存
    output_path = r"./output/test_detailed_result.docx"
    try:
        modifier.save(output_path)
        print(f"\n✅ 文件已保存: {output_path}")
        
        # 尝试重新打开验证
        print("\n🔍 验证文件完整性...")
        test_doc = Document(output_path)
        print(f"   段落数: {len(test_doc.paragraphs)}")
        print(f"   ✅ 文件可以正常打开")
        
    except Exception as e:
        print(f"\n❌ 保存或打开文件失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_modifications()
