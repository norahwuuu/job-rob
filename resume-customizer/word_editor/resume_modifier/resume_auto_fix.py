"""
生成后的简历自动修复：删除重复章节标题、为要点补全冒号前导语加粗。
"""

from __future__ import annotations

from typing import Any, Dict, List

from docx import Document
from docx.oxml.ns import qn

from .content_modifier import ContentModifier
from .resume_doc_checks import (
    BULLET_PREFIX_REST,
    iter_all_paragraphs,
    lead_before_colon_needs_marker,
    normalize_section_title,
    paragraph_is_word_list_item,
    paragraph_needs_leading_bold_fix,
    region_chars_are_bold,
)
from .resume_parser import ResumeParser


def _clear_paragraph_runs(paragraph: Any) -> None:
    el = paragraph._element
    for child in list(el):
        if child.tag == qn("w:r"):
            el.remove(child)


def _fill_paragraph_markdown(paragraph: Any, text: str) -> None:
    cm = ContentModifier.__new__(ContentModifier)
    segments = ContentModifier._parse_formatted_text(cm, text)
    for content, is_bold, is_italic in segments:
        if not content:
            continue
        run = paragraph.add_run(content)
        run.bold = is_bold
        run.italic = is_italic


def _fix_single_line_bullet_paragraph(paragraph: Any, stripped: str) -> bool:
    raw = paragraph.text
    off = len(raw) - len(raw.lstrip())
    m = BULLET_PREFIX_REST.match(stripped)
    if not m:
        if not paragraph_is_word_list_item(paragraph):
            return False
        if ":" not in stripped[:120]:
            return False
        colon_stripped = stripped.index(":", 0, 120)
        if "**" in stripped[:colon_stripped]:
            return False
        if region_chars_are_bold(paragraph, off, off + colon_stripped):
            return False
        lead = stripped[:colon_stripped].strip()
        if not lead:
            return False
        new_text = f"**{lead}**{stripped[colon_stripped:]}"
        _clear_paragraph_runs(paragraph)
        _fill_paragraph_markdown(paragraph, new_text)
        return True
    prefix, rest = m.group(1), m.group(2)
    if ":" not in rest[:120]:
        return False
    colon_idx = rest.index(":", 0, 120)
    head = rest[:colon_idx]
    if "**" in head:
        return False
    lead_start_stripped = m.end(1)
    colon_stripped = stripped.index(":", 0, 120)
    if region_chars_are_bold(
        paragraph, off + lead_start_stripped, off + colon_stripped
    ):
        return False
    lead = head.strip()
    if not lead:
        return False
    new_text = f"{prefix}**{lead}**{rest[colon_idx:]}"
    _clear_paragraph_runs(paragraph)
    _fill_paragraph_markdown(paragraph, new_text)
    return True


def _fix_multiline_bullet_paragraph(paragraph: Any) -> bool:
    """多行段落：逐行改写后按 markdown 写入（保留段落样式）。"""
    stripped = paragraph.text.strip()
    is_list = paragraph_is_word_list_item(paragraph)
    new_lines: List[str] = []
    changed = False
    for raw in stripped.splitlines():
        line = raw.strip()
        if not line:
            new_lines.append("")
            continue
        line_rule = is_list or bool(BULLET_PREFIX_REST.match(line))
        if not line_rule or not lead_before_colon_needs_marker(line):
            new_lines.append(line)
            continue
        m = BULLET_PREFIX_REST.match(line)
        if m:
            prefix, rest = m.group(1), m.group(2)
            if ":" not in rest[:120]:
                new_lines.append(line)
                continue
            colon_ij = rest.index(":", 0, 120)
            lead = rest[:colon_ij].strip()
            if not lead or "**" in rest[:colon_ij]:
                new_lines.append(line)
                continue
            new_lines.append(f"{prefix}**{lead}**{rest[colon_ij:]}")
            changed = True
            continue
        if ":" not in line[:120]:
            new_lines.append(line)
            continue
        ci = line.index(":", 0, 120)
        lead2 = line[:ci].strip()
        if not lead2 or "**" in line[:ci]:
            new_lines.append(line)
            continue
        new_lines.append(f"**{lead2}**{line[ci:]}")
        changed = True
    if not changed:
        return False
    _clear_paragraph_runs(paragraph)
    _fill_paragraph_markdown(paragraph, "\n".join(new_lines))
    return True


def _fix_bullet_paragraph(paragraph: Any) -> bool:
    stripped = paragraph.text.strip()
    if not stripped:
        return False
    if "\n" in stripped:
        return _fix_multiline_bullet_paragraph(paragraph)
    return _fix_single_line_bullet_paragraph(paragraph, stripped)


def _remove_duplicate_section_paragraphs(doc: Document) -> int:
    seen: set[str] = set()
    to_drop: List[Any] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if ResumeParser.classify_section_heading(t) is None:
            continue
        key = normalize_section_title(t)
        if key in seen:
            to_drop.append(para._element)
        else:
            seen.add(key)
    for el in to_drop:
        el.getparent().remove(el)
    return len(to_drop)


def _fix_all_bullet_leads(doc: Document, max_fixes: int = 200) -> int:
    n = 0
    for para in iter_all_paragraphs(doc):
        if n >= max_fixes:
            break
        try:
            txt = para.text
            if ":" not in txt or not txt.strip():
                continue
            if paragraph_needs_leading_bold_fix(para) and _fix_bullet_paragraph(para):
                n += 1
        except (ValueError, IndexError):
            continue
    return n


def apply_resume_auto_fixes(doc_path: str, verbose: bool = False) -> Dict[str, Any]:
    """就地修改 docx：去掉重复的章节标题段，修正带冒号要点的前导语加粗。"""
    doc = Document(doc_path)
    report: Dict[str, Any] = {
        "removed_duplicate_headings": _remove_duplicate_section_paragraphs(doc),
        "bullet_lead_fixes": 0,
        "path": doc_path,
    }
    report["bullet_lead_fixes"] = _fix_all_bullet_leads(doc)
    doc.save(doc_path)
    if verbose and (report["removed_duplicate_headings"] or report["bullet_lead_fixes"]):
        from rich.console import Console

        Console().print(
            f"  [green]✓ 自动修复：删除重复标题 {report['removed_duplicate_headings']} 处，"
            f"要点格式 {report['bullet_lead_fixes']} 处[/green]"
        )
    return report
