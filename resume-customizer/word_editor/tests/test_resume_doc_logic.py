"""单元测试：简历校验 / 自动修复 / JD 关键词（无需 cv.docx）。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from resume_modifier.ai_analyzer import extract_jd_keyword_checklist
from resume_modifier.content_modifier import ContentModifier
from resume_modifier.resume_auto_fix import apply_resume_auto_fixes
from resume_modifier.resume_doc_checks import (
    line_needs_bold_fix_text,
    normalize_section_title,
)
from resume_modifier.resume_parser import ResumeParser
from resume_modifier.resume_validator import validate_resume_docx


def test_normalize_section_title() -> None:
    assert normalize_section_title("  Work   Experience  ") == "work experience"


def test_line_needs_bold_fix_text() -> None:
    assert line_needs_bold_fix_text("• Lead: did things") is True
    assert line_needs_bold_fix_text("• **Lead:** did things") is False
    assert line_needs_bold_fix_text("plain line") is False


def test_classify_section_heading() -> None:
    assert ResumeParser.classify_section_heading("Work Experience") == "experience"
    assert ResumeParser.classify_section_heading("EDUCATION") == "education"
    assert ResumeParser.classify_section_heading("Some long paragraph " * 5) is None


def test_list_replacement_body_lines_splits_bullets() -> None:
    cm = ContentModifier.__new__(ContentModifier)
    lines = cm._list_replacement_body_lines("• **A:** one\n• **B:** two")
    assert len(lines) == 2
    assert "**A:**" in lines[0]
    assert "**B:**" in lines[1]


def test_extract_jd_keyword_checklist() -> None:
    jd = "Need Python, FastAPI, Kubernetes, CI/CD, and RAG experience."
    out = extract_jd_keyword_checklist(jd)
    assert "Python" in out or "python" in out.lower()
    assert "RAG" in out or "rag" in out.lower()


def test_validate_and_autofix_docx(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Work Experience")
    doc.add_paragraph("Work Experience")
    doc.add_paragraph("• Foo: bar description")

    path = tmp_path / "t.docx"
    doc.save(str(path))

    v0 = validate_resume_docx(str(path))
    assert v0["ok"] is False
    assert v0["duplicate_section_titles"]

    fix = apply_resume_auto_fixes(str(path), verbose=False)
    assert fix["removed_duplicate_headings"] >= 1
    assert fix["bullet_lead_fixes"] >= 1

    v1 = validate_resume_docx(str(path))
    assert v1["ok"] is True


def test_skills_line_without_list_not_flagged(tmp_path: Path) -> None:
    """非列表、无字面要点符的「Skills: …」不参与冒号前导语规则（贴合 cv.docx 混排）。"""
    doc = Document()
    doc.add_paragraph("Skills: Python, Java, Go")
    path = tmp_path / "plain.docx"
    doc.save(str(path))
    assert validate_resume_docx(str(path))["ok"] is True


def test_validate_clean_docx(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Summary")
    doc.add_paragraph("• **Thing:** value")
    path = tmp_path / "clean.docx"
    doc.save(str(path))
    v = validate_resume_docx(str(path))
    assert v["ok"] is True
