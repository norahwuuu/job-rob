from docx import Document

doc = Document('cv.docx')

print("=== 检查段落结构 ===\n")
for i, para in enumerate(doc.paragraphs[:25]):
    if 'Z. AI' in para.text or 'Beijing' in para.text or 'DataGrand' in para.text:
        print(f"段落 {i}:")
        print(f"  Text: '{para.text}'")
        print(f"  Runs count: {len(para.runs)}")
        print(f"  Runs: {[r.text for r in para.runs]}")
        print(f"  Runs sum: '{''.join([r.text for r in para.runs])}'")
        print()

print("\n=== 检查表格 ===\n")
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para_idx, para in enumerate(cell.paragraphs):
                if 'Z. AI' in para.text or 'Beijing' in para.text or 'DataGrand' in para.text:
                    print(f"表格 {table_idx}, 行 {row_idx}, 单元格 {cell_idx}, 段落 {para_idx}:")
                    print(f"  Cell text: '{cell.text[:80]}'")
                    print(f"  Para text: '{para.text}'")
                    print(f"  Runs count: {len(para.runs)}")
                    print(f"  Runs: {[r.text for r in para.runs]}")
                    print(f"  Runs sum: '{''.join([r.text for r in para.runs])}'")
                    print()
