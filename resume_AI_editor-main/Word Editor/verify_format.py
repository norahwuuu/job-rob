"""
验证格式标记是否正确应用
检查技能列表中只有类别标签和选中的术语加粗
"""
import sys
from pathlib import Path
from docx import Document
import time

sys.path.insert(0, str(Path(__file__).parent))

from resume_modifier.ai_analyzer import AIAnalyzer
from resume_modifier.content_modifier import ContentModifier

def verify_format():
    """验证格式是否正确"""
    
    print("=" * 80)
    print("🔍 验证格式标记应用")
    print("=" * 80)
    
    resume_path = r"./cv.docx"
    
    # 读取简历
    doc = Document(resume_path)
    resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    # 测试JD
    job_description = """
    AI Engineer Position - LLM Integration Focus
    
    Requirements:
    - Strong Python and FastAPI experience
    - LLM API integrations
    - RAG systems and LangChain
    - Customer collaboration
    """
    
    print("\n🤖 AI 分析...")
    analyzer = AIAnalyzer()
    result = analyzer.analyze(job_description, resume_text)
    
    print(f"   生成了 {len(result.modifications)} 条修改指令\n")
    
    # 应用修改
    modifier = ContentModifier(resume_path)
    
    for mod in result.modifications:
        modifier._apply_single_modification(mod)
    
    # 保存
    output_path = rf"./output/format_verify_{int(time.time())}.docx"
    modifier.save(output_path)
    print(f"✅ 已保存到: {output_path}\n")
    
    # 验证格式
    print("=" * 80)
    print("📊 格式验证结果")
    print("=" * 80)
    
    output_doc = Document(output_path)
    
    for para in output_doc.paragraphs:
        text = para.text
        
        # 检查技能列表行
        if any(label in text for label in ["AI & GenAI:", "Backend & Languages:", "Cloud & DevOps:", "Languages:"]):
            print(f"\n段落: {text[:60]}...")
            
            # 分析每个run的格式
            runs_info = []
            for run in para.runs:
                if run.text.strip():
                    runs_info.append({
                        'text': run.text[:30],
                        'bold': run.bold,
                        'length': len(run.text)
                    })
            
            print(f"  总共 {len(runs_info)} 个 runs:")
            for i, info in enumerate(runs_info[:10]):  # 只显示前10个
                bold_marker = "✅ BOLD" if info['bold'] else "   normal"
                print(f"    Run {i+1}: {bold_marker} | '{info['text']}'")
            
            # 检查问题
            all_bold = all(r['bold'] for r in runs_info if r['text'].strip())
            category_bold = runs_info[0]['bold'] if runs_info else False
            
            if all_bold:
                print(f"  ❌ 问题: 整行都是粗体！")
            elif category_bold:
                print(f"  ✅ 正确: 类别标签加粗，其他选择性加粗")
            else:
                print(f"  ⚠️  警告: 类别标签未加粗")
    
    print("\n" + "=" * 80)
    print(f"请打开文件验证: {output_path}")
    print("=" * 80)

if __name__ == "__main__":
    verify_format()
