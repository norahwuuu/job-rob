"""
完整测试：AI分析 + 格式标记应用
测试 prompt 更新后是否正确生成和应用粗体格式
"""
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent))

from resume_modifier.ai_analyzer import AIAnalyzer
from resume_modifier.content_modifier import ContentModifier

def test_full_pipeline():
    """测试完整流程：AI分析 + 应用修改 + 检查格式"""
    
    print("=" * 80)
    print("🧪 完整格式测试：AI 分析 + 格式标记应用")
    print("=" * 80)
    
    # 简历路径
    resume_path = r"./cv.docx"
    
    # 1. 读取简历
    print("\n📖 步骤 1: 读取简历...")
    doc = Document(resume_path)
    resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    print(f"   简历长度: {len(resume_text)} 字符")
    
    # 2. 模拟岗位描述（简单测试）
    job_description = """
    AI Engineer Position
    
    We are looking for an AI Engineer with experience in:
    - LLM API integrations and prompt engineering
    - RAG systems and semantic search
    - FastAPI backend development
    - Customer collaboration and prototyping
    - Evaluation and monitoring tools (e.g., Langfuse)
    
    The role requires strong Python skills and experience with LangChain.
    """
    
    print("\n🤖 步骤 2: AI 分析岗位描述...")
    analyzer = AIAnalyzer()
    
    try:
        result = analyzer.analyze(job_description, resume_text)
        print(f"   公司: {result.company_name}")
        print(f"   职位: {result.job_title}")
        print(f"   匹配分数: {result.match_score}")
        print(f"   修改指令数量: {len(result.modifications)}")
        
        # 3. 显示生成的修改指令（检查是否有粗体标记）
        print("\n📝 步骤 3: 检查生成的修改指令...")
        for i, mod in enumerate(result.modifications[:3], 1):  # 只显示前3条
            print(f"\n   指令 {i}:")
            print(f"     Target: {mod.target[:60]}...")
            print(f"     Replacement: {mod.replacement[:100]}...")
            print(f"     Match Type: {mod.match_type}")
            
            # 检查是否包含格式标记
            has_bold = "**" in mod.replacement
            has_bullet = mod.replacement.strip().startswith("•")
            print(f"     ✓ 包含粗体标记: {'是' if has_bold else '否'}")
            print(f"     ✓ 包含项目符号: {'是' if has_bullet else '否'}")
        
        # 4. 应用修改
        print("\n✏️ 步骤 4: 应用修改到简历...")
        modifier = ContentModifier(resume_path)
        
        success_count = 0
        for mod in result.modifications:
            if modifier._apply_single_modification(mod):
                success_count += 1
        
        print(f"   成功: {success_count}/{len(result.modifications)}")
        
        # 5. 保存并检查结果
        output_path = r"./output/test_full_format_result.docx"
        modifier.save(output_path)
        print(f"\n💾 步骤 5: 已保存到 {output_path}")
        
        # 6. 验证格式（读取保存的文档）
        print("\n🔍 步骤 6: 验证格式标记是否正确应用...")
        output_doc = Document(output_path)
        
        format_checks = []
        for para in output_doc.paragraphs:
            text = para.text
            # 检查是否还有未处理的星号
            if "**" in text:
                format_checks.append(f"❌ 发现未处理的星号: {text[:80]}")
            
            # 检查是否有粗体run
            for run in para.runs:
                if run.bold and run.text.strip():
                    format_checks.append(f"✅ 粗体文本: '{run.text.strip()}'")
                    break
        
        if format_checks:
            print("\n   格式检查结果（前10条）:")
            for check in format_checks[:10]:
                print(f"     {check}")
        
        print("\n" + "=" * 80)
        print("✅ 测试完成！")
        print("=" * 80)
        print("\n请打开文件检查：")
        print(f"  {output_path}")
        print("\n检查要点：")
        print("  1. 星号 ** 是否全部消失？")
        print("  2. 对应文本是否正确加粗？")
        print("  3. Bullet 点的前导短语（如 'Graph RAG Development:'）是否加粗？")
        print("  4. 关键技术术语（如 'LangChain', 'FastAPI'）是否加粗？")
        
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_full_pipeline()
