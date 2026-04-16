"""
简历解析器 - 读取 Word 文档并提取结构化内容
"""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
from pathlib import Path


@dataclass
class TextBlock:
    """文本块，表示一个段落或表格单元格的内容"""
    text: str
    style: Optional[str] = None
    location: str = "body"  # body, header, footer, table
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ResumeSection:
    """简历的一个部分（如工作经历、教育背景等）"""
    title: str
    content: List[TextBlock]
    start_index: int
    end_index: int


@dataclass
class ParsedResume:
    """解析后的简历结构"""
    file_path: str
    full_text: str
    sections: List[ResumeSection]
    all_blocks: List[TextBlock]
    tables: List[List[List[str]]]  # 表格数据
    

class ResumeParser:
    """
    简历解析器
    
    功能：
    1. 读取 Word 文档
    2. 提取所有文本内容（正文、表格、页眉页脚）
    3. 识别简历的各个部分
    """
    
    # 常见的简历段落标题
    SECTION_KEYWORDS = {
        'experience': ['工作经历', '工作经验', 'Work Experience', 'Experience', 'Employment', 'Professional Experience'],
        'education': ['教育背景', '教育经历', 'Education', 'Academic Background'],
        'skills': ['技能', '专业技能', 'Skills', 'Technical Skills', 'Core Competencies'],
        'projects': ['项目经历', '项目经验', 'Projects', 'Project Experience'],
        'summary': ['个人简介', '简介', 'Summary', 'Profile', 'About Me', 'Professional Summary'],
        'contact': ['联系方式', 'Contact', 'Contact Information'],
        'certifications': ['证书', '资格认证', 'Certifications', 'Certificates'],
        'languages': ['语言能力', '语言', 'Languages'],
        'awards': ['荣誉奖项', '获奖经历', 'Awards', 'Honors'],
    }
    
    def __init__(self, doc_path: str):
        """
        初始化解析器
        
        Args:
            doc_path: Word 文档路径
        """
        self.doc_path = Path(doc_path)
        if not self.doc_path.exists():
            raise FileNotFoundError(f"文件不存在: {doc_path}")
        
        self.doc = Document(doc_path)
        self.all_blocks: List[TextBlock] = []
        self.tables_data: List[List[List[str]]] = []
    
    def parse(self) -> ParsedResume:
        """
        解析文档，返回结构化的简历数据
        """
        # 1. 提取所有文本块
        self._extract_body_paragraphs()
        self._extract_tables()
        self._extract_headers_footers()
        
        # 2. 生成完整文本
        full_text = self._get_full_text()
        
        # 3. 识别简历段落
        sections = self._identify_sections()
        
        return ParsedResume(
            file_path=str(self.doc_path),
            full_text=full_text,
            sections=sections,
            all_blocks=self.all_blocks,
            tables=self.tables_data
        )
    
    def _extract_body_paragraphs(self):
        """提取正文段落"""
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip():
                self.all_blocks.append(TextBlock(
                    text=para.text,
                    style=para.style.name if para.style else None,
                    location="body",
                    metadata={"paragraph_index": i}
                ))
    
    def _extract_tables(self):
        """提取表格内容"""
        for table_idx, table in enumerate(self.doc.tables):
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    
                    if cell_text:
                        self.all_blocks.append(TextBlock(
                            text=cell_text,
                            location="table",
                            metadata={
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "cell_index": cell_idx
                            }
                        ))
                table_data.append(row_data)
            self.tables_data.append(table_data)
    
    def _extract_headers_footers(self):
        """提取页眉页脚"""
        for section_idx, section in enumerate(self.doc.sections):
            # 页眉
            for para in section.header.paragraphs:
                if para.text.strip():
                    self.all_blocks.append(TextBlock(
                        text=para.text,
                        style=para.style.name if para.style else None,
                        location="header",
                        metadata={"section_index": section_idx}
                    ))
            
            # 页脚
            for para in section.footer.paragraphs:
                if para.text.strip():
                    self.all_blocks.append(TextBlock(
                        text=para.text,
                        style=para.style.name if para.style else None,
                        location="footer",
                        metadata={"section_index": section_idx}
                    ))
    
    def _get_full_text(self) -> str:
        """获取完整文本"""
        texts = []
        for block in self.all_blocks:
            texts.append(block.text)
        return "\n".join(texts)
    
    def _identify_sections(self) -> List[ResumeSection]:
        """识别简历的各个部分"""
        sections = []
        current_section: Optional[ResumeSection] = None
        
        body_blocks = [b for b in self.all_blocks if b.location == "body"]
        
        for i, block in enumerate(body_blocks):
            section_type = self._detect_section_type(block.text)
            
            if section_type:
                # 保存当前段落
                if current_section:
                    current_section.end_index = i - 1
                    sections.append(current_section)
                
                # 开始新段落
                current_section = ResumeSection(
                    title=block.text,
                    content=[block],
                    start_index=i,
                    end_index=i
                )
            elif current_section:
                current_section.content.append(block)
        
        # 保存最后一个段落
        if current_section:
            current_section.end_index = len(body_blocks) - 1
            sections.append(current_section)
        
        return sections
    
    def _detect_section_type(self, text: str) -> Optional[str]:
        """检测文本是否是段落标题"""
        return self.classify_section_heading(text)

    @staticmethod
    def classify_section_heading(text: str) -> Optional[str]:
        """根据内置关键词判断是否为章节标题；是则返回类别键，否则 None。"""
        text_lower = text.lower().strip()
        for section_type, keywords in ResumeParser.SECTION_KEYWORDS.items():
            for keyword in keywords:
                if keyword.lower() in text_lower and len(text) < 50:
                    return section_type
        return None
    
    def get_text_for_ai(self) -> str:
        """
        获取适合发送给 AI 的文本格式
        包含位置标记，方便 AI 理解结构
        """
        lines = []
        lines.append("=== 简历内容 ===\n")
        
        for block in self.all_blocks:
            location_tag = f"[{block.location.upper()}]"
            lines.append(f"{location_tag} {block.text}")
        
        return "\n".join(lines)


def parse_resume(doc_path: str) -> ParsedResume:
    """
    便捷函数：解析简历
    
    Args:
        doc_path: Word 文档路径
        
    Returns:
        ParsedResume 对象
    """
    parser = ResumeParser(doc_path)
    return parser.parse()


if __name__ == "__main__":
    # 测试代码
    import sys
    
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        doc_path = "../cv.docx"
    
    try:
        resume = parse_resume(doc_path)
        print(f"文件: {resume.file_path}")
        print(f"总文本块数: {len(resume.all_blocks)}")
        print(f"表格数: {len(resume.tables)}")
        print(f"识别到的段落数: {len(resume.sections)}")
        print("\n段落列表:")
        for section in resume.sections:
            print(f"  - {section.title}")
        print("\n完整文本预览 (前500字符):")
        print(resume.full_text[:500])
    except Exception as e:
        print(f"解析失败: {e}")
