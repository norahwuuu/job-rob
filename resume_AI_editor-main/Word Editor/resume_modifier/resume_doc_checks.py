"""
简历 Word 文档的公共检查工具：章节标题规范化、要点前导语加粗规则、遍历段落等。
供 resume_validator 与 resume_auto_fix 共用，避免循环依赖。
"""

from __future__ import annotations

import re
from typing import Any, Iterator

from docx.oxml.ns import qn

BULLET_START = re.compile(r"^[•\-\u2022\*]\s*")
BULLET_PREFIX_REST = re.compile(r"^([•\-\u2022\*]\s*)(.+)$", re.DOTALL)


def normalize_section_title(title: str) -> str:
    return re.sub(r"\s+", " ", title.strip().lower())


def paragraph_is_word_list_item(paragraph: Any) -> bool:
    """是否为 Word 列表段（cv.docx 等模板里经历要点多为 numPr，正文无字面量 •）。"""
    try:
        p_pr = paragraph._element.pPr
        if p_pr is None:
            return False
        return p_pr.find(qn("w:numPr")) is not None
    except (AttributeError, TypeError):
        return False


def lead_before_colon_needs_marker(line: str) -> bool:
    """首段至第一个冒号内是否缺少 **（表示前导语应加粗标记）。"""
    s = line.strip()
    if ":" not in s[:120]:
        return False
    colon_idx = s.index(":", 0, 120)
    return "**" not in s[:colon_idx]


def paragraph_applies_leading_bold_rule(paragraph: Any, stripped: str) -> bool:
    """仅对 Word 列表段或带字面要点符的行做「冒号前导语」检查，避免误伤普通句子。"""
    if paragraph_is_word_list_item(paragraph):
        return True
    return bool(BULLET_START.match(stripped))


def region_chars_are_bold(paragraph: Any, start: int, end_exclusive: int) -> bool:
    """判断 paragraph.text 在 [start, end_exclusive) 内的字符是否均落在 bold run 中。"""
    if end_exclusive <= start:
        return True
    pos = 0
    for run in paragraph.runs:
        seg = run.text
        if not seg:
            continue
        run_end = pos + len(seg)
        lo = max(pos, start)
        hi = min(run_end, end_exclusive)
        if lo < hi and run.bold is not True:
            return False
        pos = run_end
    return pos >= end_exclusive


def line_needs_bold_fix_text(line: str) -> bool:
    """兼容旧调用：字面要点行且冒号前有加粗要求。"""
    if not BULLET_START.match(line):
        return False
    return lead_before_colon_needs_marker(line)


def paragraph_needs_leading_bold_fix(paragraph: Any) -> bool:
    raw = paragraph.text
    stripped = raw.strip()
    if not stripped:
        return False
    if not paragraph_applies_leading_bold_rule(paragraph, stripped):
        return False
    if "\n" in stripped:
        is_list = paragraph_is_word_list_item(paragraph)
        for ln in stripped.splitlines():
            s = ln.strip()
            if not s or ":" not in s[:120]:
                continue
            if (is_list or bool(BULLET_START.match(s))) and lead_before_colon_needs_marker(
                s
            ):
                return True
        return False
    if not lead_before_colon_needs_marker(stripped):
        return False
    colon_idx = stripped.index(":", 0, 120)
    m = BULLET_START.match(stripped)
    lead_start = m.end() if m else 0
    off = len(raw) - len(raw.lstrip())
    return not region_chars_are_bold(
        paragraph, off + lead_start, off + colon_idx
    )


def iter_all_paragraphs(doc: Any) -> Iterator[Any]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
