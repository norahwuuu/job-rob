"""
内容修改器 - 根据 AI 指令精确修改 Word 文档

增强版：支持模糊匹配、验证机制、详细错误报告
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Tuple
from pathlib import Path
import copy
import re
import unicodedata
from difflib import SequenceMatcher

from .resume_doc_checks import paragraph_is_word_list_item


@dataclass
class ModificationInstruction:
    """Modification instruction (consistent with ai_analyzer definition)"""
    target: str
    replacement: str
    reason: str
    priority: str = "medium"
    match_type: str = "fuzzy"  # exact, contains, regex, fuzzy, add_after, replace_paragraph


def normalize_text(text: str) -> str:
    """
    标准化文本，用于模糊匹配
    - 统一Unicode字符（全角->半角，中文标点->英文标点）
    - 合并多个空白字符为单个空格（包括制表符、换行符）
    - 移除首尾空白
    """
    if not text:
        return ""
    
    # Unicode 标准化
    text = unicodedata.normalize('NFKC', text)
    
    # 中文标点转英文标点
    punctuation_map = {
        '，': ',', '。': '.', '：': ':', '；': ';',
        '"': '"', '"': '"', ''': "'", ''': "'",
        '（': '(', '）': ')', '【': '[', '】': ']',
        '！': '!', '？': '?', '、': ',',
        '—': '-', '–': '-', '―': '-',
    }
    for ch, en in punctuation_map.items():
        text = text.replace(ch, en)
    
    # 合并多个空白字符（包括制表符、换行、多个空格）为单个空格
    text = re.sub(r'[\s\t\n\r]+', ' ', text)
    
    # 移除首尾空白
    text = text.strip()
    
    return text


def find_best_match(text: str, target: str, threshold: float = 0.8) -> Optional[Tuple[int, int, float]]:
    """
    在文本中找到与target最匹配的子串
    
    优化策略：
    1. 首先尝试精确匹配（最快）
    2. 然后尝试忽略大小写匹配
    3. 最后只在文本足够短时才用滑动窗口
    
    Returns:
        (start_idx, end_idx, similarity) 或 None
    """
    if not text or not target:
        return None
    
    # 1. 精确匹配（最快）
    idx = text.find(target)
    if idx >= 0:
        return (idx, idx + len(target), 1.0)
    
    # 2. 忽略大小写匹配
    text_lower = text.lower()
    target_lower = target.lower()
    idx = text_lower.find(target_lower)
    if idx >= 0:
        return (idx, idx + len(target), 0.95)
    
    # 3. 标准化后精确匹配
    normalized_target = normalize_text(target)
    normalized_text = normalize_text(text)
    
    idx = normalized_text.find(normalized_target)
    if idx >= 0:
        original_pos = _map_to_original_position(text, idx, len(normalized_target))
        if original_pos:
            return (original_pos[0], original_pos[1], 0.9)
    
    # 4. 只有在文本较短时才使用滑动窗口（避免性能问题）
    # 对于长文本，跳过滑动窗口
    if len(normalized_text) > 700 or len(normalized_target) > 100:
        return None
    
    # 滑动窗口查找最佳匹配（限制范围避免太慢）
    best_match = None
    best_similarity = threshold
    
    target_len = len(normalized_target)
    min_window = max(1, int(target_len * 0.8))  # 更窄的范围
    max_window = int(target_len * 1.2)
    
    # 限制最大迭代次数
    iterations = 0
    max_iterations = 1000
    
    for window_size in range(min_window, min(max_window + 1, len(normalized_text) + 1)):
        for i in range(len(normalized_text) - window_size + 1):
            iterations += 1
            if iterations > max_iterations:
                break
                
            window = normalized_text[i:i + window_size]
            similarity = SequenceMatcher(None, normalized_target, window).ratio()
            
            if similarity > best_similarity:
                best_similarity = similarity
                original_pos = _map_to_original_position(text, i, window_size)
                if original_pos:
                    best_match = (original_pos[0], original_pos[1], similarity)
        
        if iterations > max_iterations:
            break
    
    return best_match
    
    return best_match


def _map_to_original_position(original_text: str, normalized_start: int, normalized_len: int) -> Optional[Tuple[int, int]]:
    """
    将标准化文本中的位置映射回原始文本位置
    
    这是一个简化的映射，假设标准化主要是空白字符的合并
    """
    normalized = normalize_text(original_text)
    
    # 如果长度相同，直接返回
    if len(normalized) == len(original_text):
        return (normalized_start, normalized_start + normalized_len)
    
    # 构建位置映射表：normalized_pos -> original_pos
    mapping = []
    norm_idx = 0
    orig_idx = 0
    
    while orig_idx < len(original_text) and norm_idx < len(normalized):
        # 跳过原始文本中的连续空白
        if original_text[orig_idx].isspace():
            # 如果标准化文本当前也是空格，则匹配
            if norm_idx < len(normalized) and normalized[norm_idx] == ' ':
                mapping.append(orig_idx)
                norm_idx += 1
                orig_idx += 1
                # 跳过原始文本中的额外空白
                while orig_idx < len(original_text) and original_text[orig_idx].isspace():
                    orig_idx += 1
            else:
                orig_idx += 1
        else:
            mapping.append(orig_idx)
            norm_idx += 1
            orig_idx += 1
    
    # 确保映射表足够长
    while len(mapping) < len(normalized):
        mapping.append(len(original_text))
    
    if normalized_start >= len(mapping):
        return None
    
    start_orig = mapping[normalized_start]
    
    # 计算结束位置
    end_norm = min(normalized_start + normalized_len, len(mapping))
    if end_norm > 0 and end_norm <= len(mapping):
        end_orig = mapping[end_norm - 1] + 1
        # 扩展到原始文本中对应字符的末尾
        while end_orig < len(original_text) and original_text[end_orig].isspace():
            end_orig += 1
    else:
        end_orig = len(original_text)
    
    return (start_orig, end_orig)


@dataclass
class RunFormat:
    """Run 的格式信息"""
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    strike: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[Pt] = None
    color_rgb: Optional[Tuple[int, int, int]] = None
    
    @classmethod
    def from_run(cls, run) -> 'RunFormat':
        """从 Run 对象提取格式"""
        color_rgb = None
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            color_rgb = (rgb[0], rgb[1], rgb[2])
        
        return cls(
            bold=run.bold,
            italic=run.italic,
            underline=run.underline,
            strike=run.font.strike,
            font_name=run.font.name,
            font_size=run.font.size,
            color_rgb=color_rgb
        )
    
    def apply_to_run(self, run):
        """将格式应用到 Run"""
        if self.bold is not None:
            run.bold = self.bold
        if self.italic is not None:
            run.italic = self.italic
        if self.underline is not None:
            run.underline = self.underline
        if self.strike is not None:
            run.font.strike = self.strike
        if self.font_name:
            run.font.name = self.font_name
            # 处理中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        if self.font_size:
            run.font.size = self.font_size
        if self.color_rgb:
            run.font.color.rgb = RGBColor(*self.color_rgb)


@dataclass
class ModificationLog:
    """修改日志"""
    target: str
    replacement: str
    reason: str
    location: str
    success: bool
    error_message: Optional[str] = None
    debug_info: Optional[Dict[str, Any]] = None  # 添加调试信息字段


class ContentModifier:
    """
    内容修改器
    
    核心功能：
    1. 在 Word 文档中精确定位目标文本
    2. 执行替换操作
    3. 保持原有格式不变
    """
    
    def __init__(self, doc_path: str):
        """
        初始化修改器
        
        Args:
            doc_path: Word 文档路径
        """
        self.doc_path = Path(doc_path)
        if not self.doc_path.exists():
            raise FileNotFoundError(f"文件不存在: {doc_path}")
        
        self.doc = Document(doc_path)
        self.logs: List[ModificationLog] = []
        self.modified_paragraphs: set = set()  # 追踪已修改的段落，避免重复修改
    
    def apply_modifications(self, instructions: List[ModificationInstruction]) -> int:
        """
        应用所有修改指令
        
        Args:
            instructions: 修改指令列表
            
        Returns:
            成功修改的数量
        """
        import sys
        success_count = 0
        
        # 清空已修改段落追踪
        self.modified_paragraphs.clear()
        
        print(f"\n🔧 开始应用 {len(instructions)} 条修改指令...", flush=True)
        
        # 预处理：验证指令并过滤掉可能有问题的指令
        validated_instructions = []
        for instruction in instructions:
            # 检查target是否太短或太长
            target_len = len(instruction.target.strip())
            if target_len < 3:
                self.logs.append(ModificationLog(
                    target=instruction.target,
                    replacement=instruction.replacement,
                    reason=instruction.reason,
                    location="预检查",
                    success=False,
                    error_message=f"Target太短（{target_len}字符），跳过以避免误匹配"
                ))
                continue
            if target_len > 800:
                self.logs.append(ModificationLog(
                    target=instruction.target,
                    replacement=instruction.replacement,
                    reason=instruction.reason,
                    location="预检查",
                    success=False,
                    error_message=f"Target太长（{target_len}字符），建议拆分成多个较短的修改指令"
                ))
                continue
            
            validated_instructions.append(instruction)
        
        # 按优先级排序：high > medium > low
        priority_order = {"high": 0, "medium": 1, "low": 2}
        sorted_instructions = sorted(
            validated_instructions, 
            key=lambda x: priority_order.get(x.priority, 1)
        )
        
        for idx, instruction in enumerate(sorted_instructions, 1):
            print(f"  [{idx}/{len(sorted_instructions)}] 处理: {instruction.target[:40]}...", end="", flush=True)
            success = self._apply_single_modification(instruction)
            if success:
                success_count += 1
                print(" ✅", flush=True)
            else:
                print(" ❌", flush=True)
        
        print(f"\n✨ 完成! 成功应用 {success_count}/{len(sorted_instructions)} 条修改\n", flush=True)
        return success_count
    
    def _handle_paragraph_operation(self, instruction: ModificationInstruction) -> bool:
        """
        Handle paragraph-level operations: add_after and replace_paragraph
        
        - add_after: Insert a new paragraph after the target, preserving formatting
        - replace_paragraph: Replace entire paragraph containing the target
        """
        target = instruction.target
        replacement = instruction.replacement
        match_type = instruction.match_type
        
        # Find the paragraph containing the target
        for i, para in enumerate(self.doc.paragraphs):
            full_text = ''.join([run.text for run in para.runs])
            
            # Check if target exists in this paragraph
            match_pos = self._find_match(full_text, target, "fuzzy")
            if match_pos is None:
                continue
            
            try:
                if match_type == "add_after":
                    # For bullet points, find the last bullet in this section
                    insert_after_index = i
                    
                    # Check if replacement is a bullet point
                    is_bullet = replacement.strip().startswith(
                        ('•', '\u2022', '-', '*', '○', '▪')
                    )
                    
                    if is_bullet:
                        # Look ahead to find the last consecutive bullet point
                        for j in range(i + 1, len(self.doc.paragraphs)):
                            next_para = self.doc.paragraphs[j]
                            next_text = next_para.text.strip()
                            
                            # Check if it's a bullet point
                            if next_text.startswith(
                                ('•', '\u2022', '-', '*', '○', '▪')
                            ):
                                insert_after_index = j
                            elif not next_text:  # Empty line
                                continue
                            else:  # Not a bullet, stop here
                                break
                    
                    # Get the paragraph to insert after
                    insert_para = self.doc.paragraphs[insert_after_index]
                    p_element = insert_para._element
                    parent = p_element.getparent()
                    
                    new_elements = self._create_paragraph_elements_from_replacement(
                        insert_para, replacement
                    )
                    insert_pos = parent.index(p_element) + 1
                    for new_para in new_elements:
                        parent.insert(insert_pos, new_para)
                        insert_pos += 1
                    
                    self.logs.append(ModificationLog(
                        target=target,
                        replacement=replacement,
                        reason=instruction.reason,
                        location=f"Added after paragraph {insert_after_index}: {insert_para.text[:30]}...",
                        success=True
                    ))
                    return True
                    
                elif match_type == "replace_paragraph":
                    body_lines = self._list_replacement_body_lines(replacement)
                    if not body_lines:
                        self.logs.append(ModificationLog(
                            target=target,
                            replacement=replacement,
                            reason=instruction.reason,
                            location=f"Paragraph: {full_text[:30]}...",
                            success=False,
                            error_message="replace_paragraph: empty or whitespace replacement",
                        ))
                        return False
                    cleaned_replacement = body_lines[0]
                    segments = self._parse_formatted_text(cleaned_replacement)
                    
                    # Find a non-bold run as format template
                    template_run = None
                    for run in para.runs:
                        if not run.bold:
                            template_run = run
                            break
                    if template_run is None and para.runs:
                        template_run = para.runs[0]
                    
                    # Clear all runs
                    for run in para.runs:
                        run.text = ""
                    
                    # Remove extra runs, keep only first one
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    
                    # Add formatted runs
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    
                    for idx, (text_content, is_bold, is_italic) in enumerate(segments):
                        if not text_content:  # Skip empty segments
                            continue
                        
                        if idx == 0 and para.runs:
                            # Use existing first run
                            run = para.runs[0]
                            run.text = text_content
                        else:
                            # Create new run
                            run = para.add_run(text_content)
                        
                        # Apply template formatting
                        if template_run:
                            run.font.name = template_run.font.name
                            run.font.size = template_run.font.size
                        
                        # Apply specific formatting
                        run.bold = is_bold
                        run.italic = is_italic
                    
                    if len(body_lines) > 1:
                        p_el = para._element
                        parent = p_el.getparent()
                        ins = parent.index(p_el)
                        for extra in body_lines[1:]:
                            ins += 1
                            new_p_el = self._create_single_paragraph_element(para, extra)
                            parent.insert(ins, new_p_el)
                    
                    self.logs.append(ModificationLog(
                        target=target,
                        replacement=replacement,
                        reason=instruction.reason,
                        location=f"Replaced paragraph: {full_text[:30]}...",
                        success=True
                    ))
                    return True
                    
            except Exception as e:
                self.logs.append(ModificationLog(
                    target=target,
                    replacement=replacement,
                    reason=instruction.reason,
                    location=f"Paragraph: {full_text[:30]}...",
                    success=False,
                    error_message=f"Operation failed: {str(e)}"
                ))
                return False
        
        # Target not found
        self.logs.append(ModificationLog(
            target=target,
            replacement=replacement,
            reason=instruction.reason,
            location="Not found",
            success=False,
            error_message=f"Target not found in document: '{target[:50]}...'"
        ))
        return False
    
    def _parse_formatted_text(self, text):
        """
        Parse text with format markers and return list of (text, bold, italic) tuples
        
        Supports:
        - **text** for bold
        - *text* for italic
        - Plain text for normal
        
        Args:
            text: Text with format markers
            
        Returns:
            List of (text_content, is_bold, is_italic) tuples
        """
        import re
        
        segments = []
        pos = 0
        
        # Pattern: **text** or *text*
        # Use negative lookbehind/lookahead to avoid matching *** or ****
        pattern = r'(\*\*(?!\*)(.+?)(?<!\*)\*\*(?!\*)|\*(?!\*)(.+?)(?<!\*)\*(?!\*))'
        
        for match in re.finditer(pattern, text):
            # Add text before match as normal
            if match.start() > pos:
                segments.append((text[pos:match.start()], False, False))
            
            # Add matched text with formatting
            if match.group(0).startswith('**'):
                # Bold text
                segments.append((match.group(2), True, False))
            else:
                # Italic text
                segments.append((match.group(3), False, True))
            
            pos = match.end()
        
        # Add remaining text as normal
        if pos < len(text):
            segments.append((text[pos:], False, False))
        
        # If no formatting found, return entire text as normal
        if not segments:
            segments.append((text, False, False))
        
        return segments
    
    @staticmethod
    def _strip_one_line_bullet_prefix(line: str) -> str:
        """去掉行首要点符号（列表符号由 Word numPr 渲染）。"""
        s = line.strip()
        if not s:
            return s
        for bullet_char in ('•', '\u2022', '-', '○', '▪'):
            if s.startswith(bullet_char):
                return s[len(bullet_char):].lstrip()
        if s.startswith('*') and not s.startswith('**'):
            return s[1:].lstrip()
        return s
    
    def _strip_literal_bullets_for_word_list_paragraph(self, paragraph, text: str) -> str:
        """
        Word 列表段已由 numPr 渲染项目符号；fuzzy 等行内替换若带入字面 •/-/*，会与列表符号叠成「双黑点」。
        在写入前按行去掉行首要点字符（与 replace_paragraph / add_after 行为一致）。
        """
        if not (text or "").strip() or not paragraph_is_word_list_item(paragraph):
            return text
        norm = text.replace("\r\n", "\n").replace("\r", "\n")
        out_lines: List[str] = []
        for raw in norm.split("\n"):
            if not raw.strip():
                out_lines.append(raw)
                continue
            out_lines.append(self._strip_one_line_bullet_prefix(raw))
        return "\n".join(out_lines)

    def _list_replacement_body_lines(self, text: str) -> List[str]:
        """
        按换行拆成多条列表正文（每行去掉 •）。避免 AI 把多条要点写进一个段落导致不换行。
        """
        if not (text or '').strip():
            return []
        normalized = text.replace('\r\n', '\n').replace('\r', '\n')
        lines_out: List[str] = []
        for raw in normalized.split('\n'):
            s = raw.strip()
            if not s:
                continue
            lines_out.append(self._strip_one_line_bullet_prefix(s))
        if not lines_out:
            lines_out.append(self._strip_one_line_bullet_prefix(text.strip()))
        return lines_out
    
    def _create_single_paragraph_element(self, source_para, cleaned_text: str):
        """
        由一行已去掉行首 • 的正文生成 w:p（复制 source 的 pPr），支持 ** *。
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        new_p = OxmlElement('w:p')
        if source_para._element.pPr is not None:
            new_p.append(copy.deepcopy(source_para._element.pPr))
        
        segments = self._parse_formatted_text(cleaned_text)
        template_run = None
        if source_para.runs:
            for run in source_para.runs:
                if not run.bold:
                    template_run = run
                    break
            if template_run is None:
                template_run = source_para.runs[0]
        
        for text_content, is_bold, is_italic in segments:
            if not text_content:
                continue
            new_r = OxmlElement('w:r')
            if template_run and template_run._element.rPr is not None:
                new_rPr = copy.deepcopy(template_run._element.rPr)
                b_element = new_rPr.find(qn('w:b'))
                if b_element is not None:
                    new_rPr.remove(b_element)
                i_element = new_rPr.find(qn('w:i'))
                if i_element is not None:
                    new_rPr.remove(i_element)
                if is_bold:
                    new_rPr.append(OxmlElement('w:b'))
                if is_italic:
                    new_rPr.append(OxmlElement('w:i'))
                new_r.append(new_rPr)
            new_t = OxmlElement('w:t')
            new_t.text = text_content
            new_t.set(qn('xml:space'), 'preserve')
            new_r.append(new_t)
            new_p.append(new_r)
        return new_p
    
    def _create_paragraph_elements_from_replacement(self, source_para, text: str) -> List[Any]:
        """每条列表项一个 w:p，共享同一段落的列表样式。"""
        lines = self._list_replacement_body_lines(text)
        return [self._create_single_paragraph_element(source_para, ln) for ln in lines]
    
    def _split_bullet_paragraph_on_newlines(self, paragraph) -> None:
        """
        fuzzy 替换后若插入块含换行，将当前段拆成多个列表段，避免黑点挤在同一行。
        """
        body_lines = self._list_replacement_body_lines(paragraph.text)
        if len(body_lines) <= 1:
            return
        segments = self._parse_formatted_text(body_lines[0])
        template_run = None
        for run in paragraph.runs:
            if not run.bold:
                template_run = run
                break
        if template_run is None and paragraph.runs:
            template_run = paragraph.runs[0]
        for run in paragraph.runs:
            run.text = ""
        while len(paragraph.runs) > 1:
            paragraph._element.remove(paragraph.runs[-1]._element)
        for idx, (text_content, is_bold, is_italic) in enumerate(segments):
            if not text_content:
                continue
            if idx == 0 and paragraph.runs:
                run = paragraph.runs[0]
                run.text = text_content
            else:
                run = paragraph.add_run(text_content)
            if template_run:
                run.font.name = template_run.font.name
                run.font.size = template_run.font.size
            run.bold = is_bold
            run.italic = is_italic
        p_el = paragraph._element
        parent = p_el.getparent()
        ins = parent.index(p_el)
        for extra in body_lines[1:]:
            ins += 1
            parent.insert(ins, self._create_single_paragraph_element(paragraph, extra))
    
    def _apply_single_modification(self, instruction: ModificationInstruction) -> bool:
        """
        Apply single modification instruction
        
        Important: Each instruction only replaces the first match
        Search order: body paragraphs -> tables -> headers/footers
        """
        found = False
        
        # Handle add_after and replace_paragraph differently
        if instruction.match_type in ["add_after", "replace_paragraph"]:
            return self._handle_paragraph_operation(instruction)
        
        # 1. Process body paragraphs - only replace first match
        for para in self.doc.paragraphs:
            if self._replace_in_paragraph(para, instruction):
                found = True
                break  # Stop immediately after first match
        
        # 2. 如果正文没找到，处理表格
        if not found:
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if self._replace_in_paragraph(para, instruction):
                                found = True
                                break
                        if found:
                            break
                    if found:
                        break
                if found:
                    break
        
        # 3. 如果表格也没找到，处理页眉页脚
        if not found:
            for section in self.doc.sections:
                # 页眉
                if section.header:
                    for para in section.header.paragraphs:
                        if self._replace_in_paragraph(para, instruction):
                            found = True
                            break
                    if found:
                        break
                # 页脚
                if not found and section.footer:
                    for para in section.footer.paragraphs:
                        if self._replace_in_paragraph(para, instruction):
                            found = True
                            break
                    if found:
                        break
                if found:
                    break
        
        if not found:
            # 查找相似文本，帮助调试
            similar_texts = self.find_similar_text(instruction.target, top_n=2)
            similar_hint = ""
            if similar_texts:
                similar_hint = f" 相似文本: {[t[:40] for t, _ in similar_texts]}"
            
            self.logs.append(ModificationLog(
                target=instruction.target,
                replacement=instruction.replacement,
                reason=instruction.reason,
                location="未找到",
                success=False,
                error_message=f"未在文档中找到目标文本: '{instruction.target[:50]}...'{similar_hint}"
            ))
        
        return found
    
    def _replace_in_paragraph(self, paragraph, instruction: ModificationInstruction) -> bool:
        """
        在段落中替换文本
        
        核心算法：
        1. 合并所有 runs 的文本进行搜索
        2. 找到目标后，确定涉及哪些 runs
        3. 在第一个 run 中完成替换，清空其他 runs
        4. 保持原有格式
        
        **重要**: 使用runs的实际文本，而不是paragraph.text（可能包含隐藏内容）
        """
        # 使用段落ID追踪是否已被修改过
        para_id = id(paragraph)
        if para_id in self.modified_paragraphs:
            return False  # 已修改过的段落不再处理
        
        # 获取runs并构建实际文本
        runs = paragraph.runs
        if not runs:
            return False
        
        # 使用runs的实际文本，而不是paragraph.text
        full_text = ''.join([run.text for run in runs])
        
        target = instruction.target
        replacement = instruction.replacement
        
        # 如果段落为空或太短，跳过
        if not full_text or len(full_text.strip()) < 2:
            return False
        
        # 根据匹配类型确定是否匹配
        match_pos = self._find_match(full_text, target, instruction.match_type)
        if match_pos is None:
            return False
        
        start_idx, end_idx = match_pos
        
        # 验证匹配位置的合理性
        if start_idx < 0 or end_idx > len(full_text) or start_idx >= end_idx:
            return False
        
        affected_runs = self._find_affected_runs(runs, start_idx, end_idx)
        
        if not affected_runs:
            return False
        
        # 执行替换
        try:
            # 记录替换前的状态（用于调试）
            before_runs = [run.text for run in runs]
            before_text = full_text
            
            # 调整replacement：只保留匹配范围内需要替换的部分
            # 如果target包含段落前面的内容，需要裁剪replacement
            matched_text = full_text[start_idx:end_idx]
            adjusted_replacement = self._adjust_replacement(target, replacement, matched_text)
            adjusted_replacement = self._strip_literal_bullets_for_word_list_paragraph(
                paragraph, adjusted_replacement
            )

            self._execute_replacement(runs, affected_runs, adjusted_replacement, start_idx, end_idx)
            if instruction.match_type == "fuzzy" and any(
                ch in adjusted_replacement for ch in ("\n", "\r")
            ):
                self._split_bullet_paragraph_on_newlines(paragraph)
            
            # 记录替换后的状态（用于调试）
            after_runs = [run.text for run in runs]
            
            # 标记此段落已被修改
            self.modified_paragraphs.add(para_id)
            
            # 验证替换结果
            new_text = paragraph.text
            debug_info = {
                "before_text": before_text[:100],
                "before_runs": before_runs,
                "after_runs": after_runs,
                "affected_count": len(affected_runs),
                "match_range": f"{start_idx}-{end_idx}",
                "new_paragraph_text": new_text[:100]
            }
            
            self.logs.append(ModificationLog(
                target=target,
                replacement=replacement,
                reason=instruction.reason,
                location=f"段落: {full_text[:30]}...",
                success=True,
                debug_info=debug_info
            ))
            return True
            
        except Exception as e:
            self.logs.append(ModificationLog(
                target=target,
                replacement=replacement,
                reason=instruction.reason,
                location=f"段落: {full_text[:30]}...",
                success=False,
                error_message=str(e)
            ))
            return False
    
    def _adjust_replacement(self, target: str, replacement: str, matched_text: str) -> str:
        """
        调整replacement，确保它只替换实际匹配到的部分
        
        例如：
        - target: "Z. AI        Beijing, China"
        - replacement: "Z. AI           Berlin, Germany" 
        - matched_text: "\t   Beijing, China" （实际段落中匹配到的部分）
        
        结果应该是: "\t   Berlin, Germany" （保持相同的前缀结构）
        
        Args:
            target: AI返回的原始target
            replacement: AI返回的原始replacement
            matched_text: 段落中实际匹配到的文本
            
        Returns:
            调整后的replacement
        """
        # 标准化比较
        norm_target = normalize_text(target)
        norm_matched = normalize_text(matched_text)
        norm_replacement = normalize_text(replacement)
        
        # 如果target和matched_text标准化后完全相同，直接返回replacement
        if norm_target == norm_matched:
            return replacement
        
        # 找出target和matched_text的重叠部分
        # target可能包含matched_text没有的前缀
        if norm_matched in norm_target:
            # 计算匹配部分在target中的起始位置
            offset = norm_target.find(norm_matched)
            
            # 如果是完美后缀匹配，计算replacement中对应的部分
            if offset >= 0:
                # 尝试在replacement中找到对应的分割点
                # 假设replacement的结构与target类似
                # 例如: target="Z. AI Beijing" -> replacement="Z. AI Munich"
                #      matched="Beijing" -> 应返回 "Munich"
                
                # 简单策略：如果matched是target的后缀，则取replacement的后缀
                if norm_target.endswith(norm_matched):
                    # 计算需要保留的replacement后缀长度
                    # 这里使用启发式：保持空白字符的比例
                    prefix_len = len(norm_target) - len(norm_matched)
                    
                    # 在replacement中跳过相应长度的前缀
                    # 但保留原matched_text中的前导空白
                    leading_spaces = len(matched_text) - len(matched_text.lstrip())
                    replacement_suffix = norm_replacement[prefix_len:] if prefix_len < len(norm_replacement) else norm_replacement
                    
                    # 重建：保留matched_text的前导空白 + replacement的核心内容
                    return matched_text[:leading_spaces] + replacement_suffix.lstrip()
        
        # 如果无法智能调整，返回原replacement（可能会有问题，但至少不会崩溃）
        return replacement
    
    def _find_match(self, text: str, target: str, match_type: str) -> Optional[Tuple[int, int]]:
        """
        查找匹配位置
        
        支持多种匹配模式：
        - exact: 精确匹配
        - contains: 包含匹配
        - regex: 正则匹配
        - fuzzy: 模糊匹配（标准化后匹配，支持相似度匹配）
        
        Returns:
            (start_idx, end_idx) 或 None
        """
        if match_type == "exact":
            idx = text.find(target)
            if idx >= 0:
                return (idx, idx + len(target))
        
        elif match_type == "contains":
            idx = text.find(target)
            if idx >= 0:
                return (idx, idx + len(target))
        
        elif match_type == "regex":
            match = re.search(target, text)
            if match:
                return (match.start(), match.end())
        
        elif match_type == "fuzzy":
            # 首先尝试精确匹配（最快）
            idx = text.find(target)
            if idx >= 0:
                return (idx, idx + len(target))
            
            # 尝试忽略大小写匹配
            text_lower = text.lower()
            target_lower = target.lower()
            idx = text_lower.find(target_lower)
            if idx >= 0:
                return (idx, idx + len(target))
            
            # 尝试标准化后匹配
            normalized_text = normalize_text(text)
            normalized_target = normalize_text(target)
            
            idx = normalized_text.find(normalized_target)
            if idx >= 0:
                # 映射回原始位置
                original_pos = _map_to_original_position(text, idx, len(normalized_target))
                if original_pos:
                    return original_pos
            
            # 只有目标较短时才使用相似度匹配（避免性能问题）
            if len(target) < 400 and len(text) < 700:
                result = find_best_match(text, target, threshold=0.85)
                if result:
                    return (result[0], result[1])
        
        return None
    
    def _find_affected_runs(
        self, 
        runs, 
        start_idx: int, 
        end_idx: int
    ) -> List[Dict[str, Any]]:
        """
        找出被目标文本影响的 runs
        
        Returns:
            包含 run 信息的列表
        """
        affected = []
        current_pos = 0
        
        for i, run in enumerate(runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            # 检查是否与目标区间有交集
            if run_end > start_idx and run_start < end_idx:
                affected.append({
                    'index': i,
                    'run': run,
                    'start': run_start,
                    'end': run_end,
                    'format': RunFormat.from_run(run)
                })
            
            current_pos = run_end
        
        return affected
    
    def _execute_replacement(
        self,
        runs,
        affected_runs: List[Dict[str, Any]],
        replacement: str,
        start_idx: int,
        end_idx: int
    ):
        """
        执行实际的替换操作
        
        策略：
        1. 解析 replacement 中的格式标记 (**bold**, *italic*)
        2. 在第一个受影响的 run 中放入替换文本（带格式）
        3. 清空其他受影响的 runs
        4. 保持原有基础格式
        """
        if not affected_runs:
            return
        
        first_info = affected_runs[0]
        first_run = first_info['run']
        first_format = first_info['format']
        
        # 计算在第一个 run 中的本地位置
        local_start = max(0, start_idx - first_info['start'])
        
        # 解析格式标记
        segments = self._parse_formatted_text(replacement)
        
        # 情况1：目标完全在第一个 run 内
        if len(affected_runs) == 1 and end_idx <= first_info['end']:
            local_end = end_idx - first_info['start']
            prefix = first_run.text[:local_start]
            suffix = first_run.text[local_end:]
            
            # 如果没有格式标记，简单替换
            if len(segments) == 1 and not segments[0][1] and not segments[0][2]:
                first_run.text = prefix + replacement + suffix
            else:
                # 有格式标记，需要拆分成多个 runs
                # 使用 paragraph 对象而不是 XML 元素
                from docx.text.paragraph import Paragraph
                para_obj = Paragraph(first_run._element.getparent(), None)
                
                # 设置第一个 run 为前缀
                first_run.text = prefix
                
                # 获取第一个 run 在段落中的位置
                first_run_index = runs.index(first_run)
                
                # 为每个格式段创建新 run（在第一个 run 之后插入）
                for i, (text, is_bold, is_italic) in enumerate(segments):
                    new_run = para_obj.add_run(text)
                    
                    # 复制基础格式
                    first_format.apply_to_run(new_run)
                    
                    # 明确设置格式标记（覆盖继承的格式）
                    new_run.bold = is_bold
                    new_run.italic = is_italic
                    
                    # 将新 run 移动到正确位置（在第一个 run 之后）
                    para_element = first_run._element.getparent()
                    new_run_element = new_run._element
                    para_element.remove(new_run_element)
                    para_element.insert(first_run_index + 1 + i, new_run_element)
                
                # 添加后缀 run
                if suffix:
                    suffix_run = para_obj.add_run(suffix)
                    first_format.apply_to_run(suffix_run)
                    
                    # 移动到正确位置
                    para_element = first_run._element.getparent()
                    suffix_element = suffix_run._element
                    para_element.remove(suffix_element)
                    para_element.insert(first_run_index + 1 + len(segments), suffix_element)
            
        # 情况2：目标跨越多个 runs
        else:
            prefix = first_run.text[:local_start]
            
            # 如果没有格式标记，简单处理
            if len(segments) == 1 and not segments[0][1] and not segments[0][2]:
                first_run.text = prefix + replacement
            else:
                # 有格式标记，需要拆分成多个 runs
                from docx.text.paragraph import Paragraph
                para_obj = Paragraph(first_run._element.getparent(), None)
                
                first_run.text = prefix
                
                # 获取第一个 run 在段落中的位置
                first_run_index = runs.index(first_run)
                
                # 为每个格式段创建新 run
                for i, (text, is_bold, is_italic) in enumerate(segments):
                    new_run = para_obj.add_run(text)
                    
                    # 复制基础格式
                    first_format.apply_to_run(new_run)
                    
                    # 明确设置格式标记（覆盖继承的格式）
                    new_run.bold = is_bold
                    new_run.italic = is_italic
                    
                    # 将新 run 移动到正确位置
                    para_element = first_run._element.getparent()
                    new_run_element = new_run._element
                    para_element.remove(new_run_element)
                    para_element.insert(first_run_index + 1 + i, new_run_element)
            
            # 处理所有中间的 runs（从索引1到倒数第二个）
            if len(affected_runs) > 2:
                for info in affected_runs[1:-1]:
                    info['run'].text = ""
            
            # 最后一个 run：保留后面部分
            if len(affected_runs) > 1:
                last_info = affected_runs[-1]
                last_run = last_info['run']
                # 计算在最后一个run中的本地结束位置
                local_end = min(end_idx - last_info['start'], len(last_run.text))
                # 保留从local_end到结尾的部分
                suffix = last_run.text[local_end:]
                last_run.text = suffix
    
    def save(self, output_path: str) -> str:
        """
        保存修改后的文档
        
        Args:
            output_path: 输出路径
            
        Returns:
            实际保存的路径
        """
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output))
        return str(output)
    
    def get_all_text(self) -> str:
        """获取文档中的所有文本，用于验证"""
        texts = []
        
        # 正文段落
        for para in self.doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        
        # 表格
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            texts.append(para.text)
        
        # 页眉页脚
        for section in self.doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)
        
        return "\n".join(texts)
    
    def verify_modification(self, instruction: ModificationInstruction) -> bool:
        """
        验证修改是否生效
        
        检查 replacement 是否存在于文档中
        """
        all_text = self.get_all_text()
        return instruction.replacement in all_text
    
    def find_similar_text(self, target: str, top_n: int = 3) -> List[Tuple[str, float]]:
        """
        在文档中查找与target最相似的文本片段
        
        用于调试：当匹配失败时，显示可能的候选项
        """
        all_text = self.get_all_text()
        normalized_target = normalize_text(target)
        target_len = len(normalized_target)
        
        candidates = []
        
        # 从文档中提取所有可能的片段
        for para in self.doc.paragraphs:
            text = para.text
            if not text.strip():
                continue
            
            normalized_text = normalize_text(text)
            
            # 滑动窗口
            window_size = min(target_len * 2, len(normalized_text))
            for i in range(max(1, len(normalized_text) - window_size + 1)):
                window = normalized_text[i:i + window_size]
                similarity = SequenceMatcher(None, normalized_target, window).ratio()
                if similarity > 0.5:  # 只保留相似度 > 50% 的
                    # 获取原始文本
                    original_window = text[i:i + window_size] if i + window_size <= len(text) else text[i:]
                    candidates.append((original_window, similarity))
        
        # 去重并排序
        seen = set()
        unique_candidates = []
        for text, sim in sorted(candidates, key=lambda x: -x[1]):
            key = text[:50]  # 用前50个字符去重
            if key not in seen:
                seen.add(key)
                unique_candidates.append((text, sim))
        
        return unique_candidates[:top_n]
    
    def get_logs(self) -> List[ModificationLog]:
        """获取修改日志"""
        return self.logs
    
    def get_modification_results(self) -> List[Dict[str, Any]]:
        """
        获取修改结果列表（用于 API 返回）
        
        Returns:
            每条修改的详细结果，包含 success 状态
        """
        return [
            {
                "target": log.target,
                "replacement": log.replacement,
                "reason": log.reason,
                "success": log.success,
                "error_message": log.error_message,
                "location": log.location
            }
            for log in self.logs
        ]
    
    def get_summary(self) -> Dict[str, Any]:
        """获取修改摘要"""
        success = [log for log in self.logs if log.success]
        failed = [log for log in self.logs if not log.success]
        
        return {
            "total": len(self.logs),
            "success": len(success),
            "failed": len(failed),
            "success_details": [
                {"target": log.target[:50], "replacement": log.replacement[:50]}
                for log in success
            ],
            "failed_details": [
                {"target": log.target[:50], "error": log.error_message}
                for log in failed
            ]
        }


def modify_resume(
    doc_path: str,
    instructions: List[ModificationInstruction],
    output_path: str
) -> Dict[str, Any]:
    """
    便捷函数：修改简历
    
    Args:
        doc_path: 原始文档路径
        instructions: 修改指令列表
        output_path: 输出路径
        
    Returns:
        修改摘要
    """
    modifier = ContentModifier(doc_path)
    modifier.apply_modifications(instructions)
    modifier.save(output_path)
    return modifier.get_summary()


if __name__ == "__main__":
    # 测试代码
    test_instructions = [
        ModificationInstruction(
            target="Berlin",
            replacement="北京",
            reason="岗位要求工作地点在中国",
            priority="high"
        ),
        ModificationInstruction(
            target="3 years",
            replacement="5 years",
            reason="岗位要求5年以上经验",
            priority="medium"
        )
    ]
    
    import sys
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        doc_path = "../cv.docx"
    
    try:
        modifier = ContentModifier(doc_path)
        count = modifier.apply_modifications(test_instructions)
        print(f"成功修改 {count} 处")
        
        summary = modifier.get_summary()
        print(f"\n修改摘要:")
        print(f"  总计: {summary['total']}")
        print(f"  成功: {summary['success']}")
        print(f"  失败: {summary['failed']}")
        
        if summary['success_details']:
            print(f"\n成功的修改:")
            for detail in summary['success_details']:
                print(f"  {detail['target']} -> {detail['replacement']}")
        
        if summary['failed_details']:
            print(f"\n失败的修改:")
            for detail in summary['failed_details']:
                print(f"  {detail['target']}: {detail['error']}")
        
        # 保存测试输出
        output_path = "./test_output.docx"
        modifier.save(output_path)
        print(f"\n已保存到: {output_path}")
        
    except Exception as e:
        print(f"测试失败: {e}")
        import traceback
        traceback.print_exc()
